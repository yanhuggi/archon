"""The public ``export_issue`` MCP tool."""

from __future__ import annotations

import json
import re
import tempfile
from pathlib import Path
from typing import Annotated

from docx import Document
from mcp.server import MCPServer
from mcp.types import ToolAnnotations
from pydantic import Field

from server.config import JiraConfig
from server.files import OutputPathError, commit_output_file, resolve_output_path
from server.instructions import EXPORT_ISSUE_DESCRIPTION
from server.tools._common import error_response, provider_or_error
from server.tools.get_issue import ISSUE_KEY_RE

MAX_EMBEDDED_TEXT_CHARS = 200_000
_INVALID_XML_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _text(value: object) -> str:
    return "" if value is None else _INVALID_XML_CHARS_RE.sub("", str(value))


def register(
    mcp: MCPServer,
    default_provider: str = "jira",
    config: JiraConfig | None = None,
) -> None:
    @mcp.tool(
        name="export_issue",
        title="Export Jira Issue",
        description=EXPORT_ISSUE_DESCRIPTION,
        annotations=ToolAnnotations(
            readOnlyHint=False,
            destructiveHint=True,
            idempotentHint=False,
            openWorldHint=True,
        ),
        structured_output=False,
    )
    def export_issue(
        issue_key: Annotated[str, Field(min_length=3, max_length=100)],
        save_to: Annotated[str, Field(min_length=1, max_length=4096, description="Authorized local .docx path.")],
        include_attachments: bool = True,
    ) -> str:
        normalized_key = issue_key.strip().upper() if isinstance(issue_key, str) else ""
        if not ISSUE_KEY_RE.fullmatch(normalized_key):
            return error_response("invalid_issue_key", f"Invalid issue key: {issue_key!r}", issue_key=normalized_key)

        runtime_config = config or JiraConfig.from_env()
        try:
            save_path = resolve_output_path(save_to, runtime_config, suffix=".docx")
        except OutputPathError as exc:
            return error_response("invalid_output_path", str(exc), issue_key=normalized_key)

        provider, error = provider_or_error(default_provider)
        if error:
            return error_response("provider_unavailable", "Jira provider is unavailable", issue_key=normalized_key)
        try:
            issue = provider.get_issue_json(normalized_key)
        except Exception as exc:  # noqa: BLE001 - provider boundary
            return error_response(
                "provider_error",
                f"Jira provider failed: {type(exc).__name__}: {exc}",
                issue_key=normalized_key,
            )
        if not isinstance(issue, dict):
            return error_response(
                "invalid_provider_response",
                "Jira provider returned an invalid issue object",
                issue_key=normalized_key,
            )
        if "error" in issue:
            issue.setdefault("issue_key", normalized_key)
            issue.setdefault("error_code", "provider_error")
            return json.dumps(issue, ensure_ascii=False)

        save_path.parent.mkdir(parents=True, exist_ok=True)
        temporary_doc: Path | None = None
        try:
            document = Document()
            document.add_heading(f"{normalized_key} - {_text(issue.get('summary'))}", 0)

            document.add_heading("Issue Details", level=1)
            table = document.add_table(rows=0, cols=2)
            table.style = "Table Grid"
            metadata = (
                ("Type", issue.get("issue_type")),
                ("Status", issue.get("status")),
                ("Priority", issue.get("priority")),
                ("Assignee", issue.get("assignee", "Unassigned")),
                ("Reporter", issue.get("reporter")),
                ("Created", issue.get("created")),
                ("Updated", issue.get("updated")),
                ("Resolution", issue.get("resolution", "Unresolved")),
            )
            for label, value in metadata:
                row = table.add_row()
                row.cells[0].text = label
                row.cells[1].text = _text(value)

            description = _text(issue.get("description"))
            if description:
                document.add_heading("Description", level=1)
                document.add_paragraph(description)

            subtasks = issue.get("subtasks") or []
            if subtasks:
                document.add_heading("Subtasks", level=1)
                for subtask in subtasks:
                    if isinstance(subtask, dict):
                        document.add_paragraph(
                            f"{_text(subtask.get('key'))} - {_text(subtask.get('summary'))} "
                            f"({_text(subtask.get('status'))})",
                            style="List Bullet",
                        )

            issue_links = issue.get("issue_links") or []
            if issue_links:
                document.add_heading("Related Issues", level=1)
                for link in issue_links:
                    if not isinstance(link, dict):
                        continue
                    linked_issue = link.get("issue") or link.get("linked_issue") or {}
                    if not isinstance(linked_issue, dict):
                        continue
                    relationship = link.get("direction") or link.get("relationship") or "Related to"
                    document.add_paragraph(
                        f"{_text(relationship)} {_text(linked_issue.get('key'))} - "
                        f"{_text(linked_issue.get('summary'))} ({_text(linked_issue.get('status'))})",
                        style="List Bullet",
                    )

            attachments = issue.get("attachments") or []
            if attachments and include_attachments:
                document.add_heading("Attachments", level=1)
                with tempfile.TemporaryDirectory(prefix=".archon-jira-", dir=save_path.parent) as temp_dir:
                    for index, attachment in enumerate(attachments):
                        if not isinstance(attachment, dict):
                            continue
                        attachment_id = _text(attachment.get("id"))
                        filename = _text(attachment.get("filename"))
                        mime_type = _text(attachment.get("mime_type"))
                        document.add_heading(filename or f"Attachment {attachment_id}", level=2)
                        document.add_paragraph(f"ID: {attachment_id}")
                        document.add_paragraph(f"Size: {int(attachment.get('size') or 0) / 1024 / 1024:.2f} MB")
                        document.add_paragraph(f"Author: {_text(attachment.get('author'))}")
                        document.add_paragraph(f"Created: {_text(attachment.get('created'))}")
                        document.add_paragraph(f"MIME Type: {mime_type}")

                        is_text = mime_type.startswith("text/") or Path(filename).suffix.lower() in {
                            ".md", ".txt", ".json", ".csv", ".xml", ".yaml", ".yml"
                        }
                        if not is_text or not attachment_id.isdigit():
                            continue
                        temp_file = Path(temp_dir) / f"attachment-{index}-{attachment_id}.txt"
                        try:
                            result = json.loads(provider.get_attachment(attachment_id, str(temp_file)))
                            if "error" not in result and temp_file.exists():
                                content = temp_file.read_text(encoding="utf-8", errors="replace")
                                if len(content) > MAX_EMBEDDED_TEXT_CHARS:
                                    content = content[:MAX_EMBEDDED_TEXT_CHARS] + "\n[truncated]"
                                document.add_paragraph("Content:")
                                document.add_paragraph(_text(content))
                            elif "error" in result:
                                document.add_paragraph(f"(Attachment content unavailable: {result['error']})")
                        except Exception as exc:  # noqa: BLE001 - attachment enrichment is optional
                            document.add_paragraph(f"(Attachment content unavailable: {type(exc).__name__}: {exc})")

            with tempfile.NamedTemporaryFile(
                prefix=".archon-jira-",
                suffix=".docx.tmp",
                dir=save_path.parent,
                delete=False,
            ) as handle:
                temporary_doc = Path(handle.name)
            document.save(str(temporary_doc))
            commit_output_file(temporary_doc, save_path, runtime_config)
            temporary_doc = None
            return json.dumps(
                {
                    "issue_key": normalized_key,
                    "saved_to": str(save_path),
                    "filename": save_path.name,
                    "size": save_path.stat().st_size,
                    "includes_attachments": include_attachments,
                    "attachment_count": len(attachments),
                },
                ensure_ascii=False,
            )
        except OutputPathError as exc:
            return error_response("invalid_output_path", str(exc), issue_key=normalized_key)
        except Exception as exc:  # noqa: BLE001 - export errors need a stable tool response
            return error_response(
                "export_failed",
                f"Export failed: {type(exc).__name__}: {exc}",
                issue_key=normalized_key,
            )
        finally:
            if temporary_doc is not None:
                temporary_doc.unlink(missing_ok=True)
